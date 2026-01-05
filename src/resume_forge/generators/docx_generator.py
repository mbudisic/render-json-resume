"""Native DOCX generator using python-docx."""

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from ..schema import Resume
from .base import BaseGenerator


class DOCXGenerator(BaseGenerator):
    """Generate native DOCX documents from JSON Resume."""
    
    STYLES = {
        "professional": {
            "primary_color": RGBColor(0x2c, 0x3e, 0x50),
            "secondary_color": RGBColor(0x7f, 0x8c, 0x8d),
            "accent_color": RGBColor(0x34, 0x98, 0xdb),
            "font_name": "Calibri",
        },
        "modern": {
            "primary_color": RGBColor(0x1a, 0x1a, 0x2e),
            "secondary_color": RGBColor(0x4a, 0x4a, 0x4a),
            "accent_color": RGBColor(0xe9, 0x45, 0x60),
            "font_name": "Arial",
        },
        "elegant": {
            "primary_color": RGBColor(0x2d, 0x34, 0x36),
            "secondary_color": RGBColor(0x63, 0x6e, 0x72),
            "accent_color": RGBColor(0x6c, 0x5c, 0xe7),
            "font_name": "Times New Roman",
        },
        "minimal": {
            "primary_color": RGBColor(0x00, 0x00, 0x00),
            "secondary_color": RGBColor(0x55, 0x55, 0x55),
            "accent_color": RGBColor(0x00, 0x00, 0x00),
            "font_name": "Calibri",
        },
    }
    
    def __init__(self, resume: Resume, style: str = "professional"):
        super().__init__(resume, style)
        self.theme = self.STYLES.get(style, self.STYLES["professional"])
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self):
        """Set up document margins and base styles."""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
    
    def generate(self, output_path: Path) -> None:
        """Generate the DOCX document."""
        if self.resume.basics:
            self._build_header()
        
        if self.resume.work:
            self._build_work_section()
        
        if self.resume.education:
            self._build_education_section()
        
        if self.resume.skills:
            self._build_skills_section()
        
        if self.resume.projects:
            self._build_projects_section()
        
        if self.resume.certificates:
            self._build_certificates_section()
        
        if self.resume.awards:
            self._build_awards_section()
        
        if self.resume.publications:
            self._build_publications_section()
        
        if self.resume.volunteer:
            self._build_volunteer_section()
        
        if self.resume.languages:
            self._build_languages_section()
        
        if self.resume.interests:
            self._build_interests_section()
        
        if self.resume.references:
            self._build_references_section()
        
        self.doc.save(str(output_path))
    
    def _add_paragraph(
        self,
        text: str,
        font_size: int = 11,
        bold: bool = False,
        italic: bool = False,
        color: RGBColor = None,
        alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
        space_after: int = 6,
        space_before: int = 0,
    ):
        """Add a formatted paragraph to the document."""
        para = self.doc.add_paragraph()
        para.alignment = alignment
        para.paragraph_format.space_after = Pt(space_after)
        para.paragraph_format.space_before = Pt(space_before)
        
        run = para.add_run(text)
        run.font.name = self.theme["font_name"]
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
        
        return para
    
    def _add_section_title(self, title: str):
        """Add a section title."""
        para = self._add_paragraph(
            title,
            font_size=12,
            bold=True,
            color=self.theme["accent_color"],
            space_before=12,
            space_after=6,
        )
        return para
    
    def _build_header(self):
        """Build the header section."""
        basics = self.resume.basics
        
        if basics.name:
            self._add_paragraph(
                basics.name,
                font_size=24,
                bold=True,
                color=self.theme["primary_color"],
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_after=4,
            )
        
        if basics.label:
            self._add_paragraph(
                basics.label,
                font_size=14,
                color=self.theme["secondary_color"],
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_after=8,
            )
        
        contact_parts = []
        if basics.email:
            contact_parts.append(basics.email)
        if basics.phone:
            contact_parts.append(basics.phone)
        if basics.url:
            contact_parts.append(basics.url)
        if basics.location:
            loc = basics.location
            loc_parts = [p for p in [loc.city, loc.region, loc.countryCode] if p]
            if loc_parts:
                contact_parts.append(", ".join(loc_parts))
        
        if contact_parts:
            self._add_paragraph(
                " | ".join(contact_parts),
                font_size=10,
                color=self.theme["secondary_color"],
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_after=8,
            )
        
        if basics.profiles:
            profile_parts = []
            for profile in basics.profiles:
                if profile.network and profile.username:
                    profile_parts.append(f"{profile.network}: {profile.username}")
                elif profile.url:
                    profile_parts.append(profile.url)
            if profile_parts:
                self._add_paragraph(
                    " | ".join(profile_parts),
                    font_size=10,
                    color=self.theme["secondary_color"],
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    space_after=12,
                )
        
        if basics.summary:
            self._add_paragraph(
                basics.summary,
                font_size=10,
                color=self.theme["primary_color"],
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_after=12,
            )
    
    def _build_work_section(self):
        """Build the work experience section."""
        self._add_section_title("EXPERIENCE")
        
        for job in self.resume.work:
            title_parts = []
            if job.position:
                title_parts.append(job.position)
            if job.name:
                title_parts.append(f"at {job.name}")
            if title_parts:
                self._add_paragraph(
                    " ".join(title_parts),
                    font_size=11,
                    bold=True,
                    color=self.theme["primary_color"],
                    space_after=2,
                )
            
            date_range = self.format_date_range(job.startDate, job.endDate)
            if date_range:
                self._add_paragraph(
                    date_range,
                    font_size=10,
                    color=self.theme["secondary_color"],
                    space_after=4,
                )
            
            if job.summary:
                self._add_paragraph(
                    job.summary,
                    font_size=10,
                    color=self.theme["primary_color"],
                    space_after=4,
                )
            
            if job.highlights:
                for highlight in job.highlights:
                    para = self._add_paragraph(
                        f"• {highlight}",
                        font_size=10,
                        color=self.theme["primary_color"],
                        space_after=2,
                    )
                    para.paragraph_format.left_indent = Inches(0.25)
    
    def _build_education_section(self):
        """Build the education section."""
        self._add_section_title("EDUCATION")
        
        for edu in self.resume.education:
            title_parts = []
            if edu.studyType:
                title_parts.append(edu.studyType)
            if edu.area:
                title_parts.append(f"in {edu.area}")
            if title_parts:
                self._add_paragraph(
                    " ".join(title_parts),
                    font_size=11,
                    bold=True,
                    color=self.theme["primary_color"],
                    space_after=2,
                )
            
            if edu.institution:
                subtitle = edu.institution
                if edu.score:
                    subtitle += f" | GPA: {edu.score}"
                self._add_paragraph(
                    subtitle,
                    font_size=10,
                    color=self.theme["secondary_color"],
                    space_after=2,
                )
            
            date_range = self.format_date_range(edu.startDate, edu.endDate)
            if date_range:
                self._add_paragraph(
                    date_range,
                    font_size=10,
                    color=self.theme["secondary_color"],
                    space_after=6,
                )
    
    def _build_skills_section(self):
        """Build the skills section."""
        self._add_section_title("SKILLS")
        
        for skill in self.resume.skills:
            if skill.name:
                para = self.doc.add_paragraph()
                para.paragraph_format.space_after = Pt(4)
                
                run = para.add_run(skill.name)
                run.font.name = self.theme["font_name"]
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = self.theme["primary_color"]
                
                if skill.level:
                    run = para.add_run(f" ({skill.level})")
                    run.font.name = self.theme["font_name"]
                    run.font.size = Pt(10)
                    run.font.color.rgb = self.theme["secondary_color"]
                
                if skill.keywords:
                    run = para.add_run(": " + ", ".join(skill.keywords))
                    run.font.name = self.theme["font_name"]
                    run.font.size = Pt(10)
                    run.font.color.rgb = self.theme["primary_color"]
    
    def _build_projects_section(self):
        """Build the projects section."""
        self._add_section_title("PROJECTS")
        
        for project in self.resume.projects:
            if project.name:
                self._add_paragraph(
                    project.name,
                    font_size=11,
                    bold=True,
                    color=self.theme["primary_color"],
                    space_after=2,
                )
            
            date_range = self.format_date_range(project.startDate, project.endDate)
            if date_range:
                self._add_paragraph(
                    date_range,
                    font_size=10,
                    color=self.theme["secondary_color"],
                    space_after=4,
                )
            
            if project.description:
                self._add_paragraph(
                    project.description,
                    font_size=10,
                    color=self.theme["primary_color"],
                    space_after=4,
                )
            
            if project.highlights:
                for highlight in project.highlights:
                    para = self._add_paragraph(
                        f"• {highlight}",
                        font_size=10,
                        color=self.theme["primary_color"],
                        space_after=2,
                    )
                    para.paragraph_format.left_indent = Inches(0.25)
    
    def _build_certificates_section(self):
        """Build the certificates section."""
        self._add_section_title("CERTIFICATES")
        
        for cert in self.resume.certificates:
            if cert.name:
                self._add_paragraph(
                    cert.name,
                    font_size=11,
                    bold=True,
                    color=self.theme["primary_color"],
                    space_after=2,
                )
            
            subtitle_parts = []
            if cert.issuer:
                subtitle_parts.append(cert.issuer)
            if cert.date:
                subtitle_parts.append(self._format_date(cert.date))
            if subtitle_parts:
                self._add_paragraph(
                    " | ".join(subtitle_parts),
                    font_size=10,
                    color=self.theme["secondary_color"],
                    space_after=6,
                )
    
    def _build_awards_section(self):
        """Build the awards section."""
        self._add_section_title("AWARDS")
        
        for award in self.resume.awards:
            if award.title:
                self._add_paragraph(
                    award.title,
                    font_size=11,
                    bold=True,
                    color=self.theme["primary_color"],
                    space_after=2,
                )
            
            subtitle_parts = []
            if award.awarder:
                subtitle_parts.append(award.awarder)
            if award.date:
                subtitle_parts.append(self._format_date(award.date))
            if subtitle_parts:
                self._add_paragraph(
                    " | ".join(subtitle_parts),
                    font_size=10,
                    color=self.theme["secondary_color"],
                    space_after=4,
                )
            
            if award.summary:
                self._add_paragraph(
                    award.summary,
                    font_size=10,
                    color=self.theme["primary_color"],
                    space_after=6,
                )
    
    def _build_publications_section(self):
        """Build the publications section."""
        self._add_section_title("PUBLICATIONS")
        
        for pub in self.resume.publications:
            if pub.name:
                self._add_paragraph(
                    pub.name,
                    font_size=11,
                    bold=True,
                    color=self.theme["primary_color"],
                    space_after=2,
                )
            
            subtitle_parts = []
            if pub.publisher:
                subtitle_parts.append(pub.publisher)
            if pub.releaseDate:
                subtitle_parts.append(self._format_date(pub.releaseDate))
            if subtitle_parts:
                self._add_paragraph(
                    " | ".join(subtitle_parts),
                    font_size=10,
                    color=self.theme["secondary_color"],
                    space_after=4,
                )
            
            if pub.summary:
                self._add_paragraph(
                    pub.summary,
                    font_size=10,
                    color=self.theme["primary_color"],
                    space_after=6,
                )
    
    def _build_volunteer_section(self):
        """Build the volunteer section."""
        self._add_section_title("VOLUNTEER")
        
        for vol in self.resume.volunteer:
            title_parts = []
            if vol.position:
                title_parts.append(vol.position)
            if vol.organization:
                title_parts.append(f"at {vol.organization}")
            if title_parts:
                self._add_paragraph(
                    " ".join(title_parts),
                    font_size=11,
                    bold=True,
                    color=self.theme["primary_color"],
                    space_after=2,
                )
            
            date_range = self.format_date_range(vol.startDate, vol.endDate)
            if date_range:
                self._add_paragraph(
                    date_range,
                    font_size=10,
                    color=self.theme["secondary_color"],
                    space_after=4,
                )
            
            if vol.summary:
                self._add_paragraph(
                    vol.summary,
                    font_size=10,
                    color=self.theme["primary_color"],
                    space_after=4,
                )
            
            if vol.highlights:
                for highlight in vol.highlights:
                    para = self._add_paragraph(
                        f"• {highlight}",
                        font_size=10,
                        color=self.theme["primary_color"],
                        space_after=2,
                    )
                    para.paragraph_format.left_indent = Inches(0.25)
    
    def _build_languages_section(self):
        """Build the languages section."""
        self._add_section_title("LANGUAGES")
        
        lang_texts = []
        for lang in self.resume.languages:
            if lang.language:
                text = lang.language
                if lang.fluency:
                    text += f" ({lang.fluency})"
                lang_texts.append(text)
        
        if lang_texts:
            self._add_paragraph(
                ", ".join(lang_texts),
                font_size=10,
                color=self.theme["primary_color"],
                space_after=6,
            )
    
    def _build_interests_section(self):
        """Build the interests section."""
        self._add_section_title("INTERESTS")
        
        for interest in self.resume.interests:
            if interest.name:
                para = self.doc.add_paragraph()
                para.paragraph_format.space_after = Pt(4)
                
                run = para.add_run(interest.name)
                run.font.name = self.theme["font_name"]
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = self.theme["primary_color"]
                
                if interest.keywords:
                    run = para.add_run(": " + ", ".join(interest.keywords))
                    run.font.name = self.theme["font_name"]
                    run.font.size = Pt(10)
                    run.font.color.rgb = self.theme["primary_color"]
    
    def _build_references_section(self):
        """Build the references section."""
        self._add_section_title("REFERENCES")
        
        for ref in self.resume.references:
            if ref.name:
                self._add_paragraph(
                    ref.name,
                    font_size=11,
                    bold=True,
                    color=self.theme["primary_color"],
                    space_after=2,
                )
            if ref.reference:
                self._add_paragraph(
                    f'"{ref.reference}"',
                    font_size=10,
                    italic=True,
                    color=self.theme["primary_color"],
                    space_after=6,
                )
